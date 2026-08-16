from docx import Document
import sys

if len(sys.argv) < 2:
    print("Usage: extract_docx_full.py <docx-path> <out-path (optional)>")
    sys.exit(1)

docx_path = sys.argv[1]
out_path = sys.argv[2] if len(sys.argv) > 2 else "results/submission_guidelines_full.txt"

doc = Document(docx_path)
with open(out_path, "w", encoding="utf-8") as f:
    # paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            f.write(text + "\n")
    # tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                f.write("\t".join(cells) + "\n")

print(f"Wrote extracted text to {out_path}")
